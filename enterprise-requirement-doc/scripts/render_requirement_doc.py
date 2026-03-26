#!/usr/bin/env python3
import argparse
import re
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_BREAK
    from docx.oxml.ns import qn
    from docx.shared import Pt
except ImportError as exc:
    raise SystemExit(
        "Missing dependency: python-docx. Install it with `pip install python-docx`."
    ) from exc


HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
ORDERED_RE = re.compile(r"^\d+\.\s+(.*)$")
UNORDERED_RE = re.compile(r"^[-*]\s+(.*)$")
TABLE_SEP_RE = re.compile(r"^\|\s*[-: ]+\|\s*$")


def set_default_font(document):
    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)


def is_table_row(line):
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|")


def parse_table_row(line):
    parts = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return parts


def add_heading(document, level, text):
    paragraph = document.add_heading(level=min(level, 4))
    paragraph.add_run(text)


def add_code_block(document, code_lines):
    for line in code_lines:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(9)


def add_table(document, rows):
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    for row_idx, row in enumerate(rows):
        padded = row + [""] * (col_count - len(row))
        for col_idx, value in enumerate(padded):
            table.cell(row_idx, col_idx).text = value


def flush_paragraph_buffer(document, buffer_lines):
    if not buffer_lines:
        return
    text = " ".join(line.strip() for line in buffer_lines if line.strip())
    if text:
        document.add_paragraph(text)


def markdown_to_docx(markdown_text, output_path):
    document = Document()
    set_default_font(document)

    lines = markdown_text.splitlines()
    idx = 0
    paragraph_buffer = []

    while idx < len(lines):
        line = lines[idx]
        stripped = line.strip()

        if not stripped:
            flush_paragraph_buffer(document, paragraph_buffer)
            paragraph_buffer = []
            idx += 1
            continue

        if stripped.startswith("```"):
            flush_paragraph_buffer(document, paragraph_buffer)
            paragraph_buffer = []
            idx += 1
            code_lines = []
            while idx < len(lines) and not lines[idx].strip().startswith("```"):
                code_lines.append(lines[idx])
                idx += 1
            add_code_block(document, code_lines)
            idx += 1
            continue

        heading_match = HEADING_RE.match(stripped)
        if heading_match:
            flush_paragraph_buffer(document, paragraph_buffer)
            paragraph_buffer = []
            level = len(heading_match.group(1))
            add_heading(document, level, heading_match.group(2).strip())
            idx += 1
            continue

        if is_table_row(stripped) and idx + 1 < len(lines) and TABLE_SEP_RE.match(lines[idx + 1].strip()):
            flush_paragraph_buffer(document, paragraph_buffer)
            paragraph_buffer = []
            rows = [parse_table_row(lines[idx])]
            idx += 2
            while idx < len(lines) and is_table_row(lines[idx].strip()):
                rows.append(parse_table_row(lines[idx]))
                idx += 1
            add_table(document, rows)
            continue

        ordered_match = ORDERED_RE.match(stripped)
        if ordered_match:
            flush_paragraph_buffer(document, paragraph_buffer)
            paragraph_buffer = []
            document.add_paragraph(ordered_match.group(1).strip(), style="List Number")
            idx += 1
            continue

        unordered_match = UNORDERED_RE.match(stripped)
        if unordered_match:
            flush_paragraph_buffer(document, paragraph_buffer)
            paragraph_buffer = []
            document.add_paragraph(unordered_match.group(1).strip(), style="List Bullet")
            idx += 1
            continue

        if stripped == "---":
            flush_paragraph_buffer(document, paragraph_buffer)
            paragraph_buffer = []
            paragraph = document.add_paragraph()
            run = paragraph.add_run()
            run.add_break(WD_BREAK.PAGE)
            idx += 1
            continue

        paragraph_buffer.append(line)
        idx += 1

    flush_paragraph_buffer(document, paragraph_buffer)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))


def main():
    parser = argparse.ArgumentParser(description="Render a markdown requirement doc to DOCX.")
    parser.add_argument("input", help="Input markdown path")
    parser.add_argument("output", help="Output docx path")
    args = parser.parse_args()

    input_path = Path(args.input)
    output_path = Path(args.output)

    if not input_path.exists():
        raise SystemExit(f"Input file does not exist: {input_path}")

    markdown_text = input_path.read_text(encoding="utf-8")
    markdown_to_docx(markdown_text, output_path)


if __name__ == "__main__":
    main()
