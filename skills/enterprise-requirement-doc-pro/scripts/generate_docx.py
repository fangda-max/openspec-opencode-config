#!/usr/bin/env python3
"""
需求文档 DOCX 生成工具 - 增强版
支持更好的格式转换
"""

import sys
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_run_font(run, font_name='微软雅黑', font_size=Pt(10.5), bold=False, italic=False, color=None):
    """设置 run 的字体样式"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = font_size
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def create_heading(doc, text, level):
    """创建标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        font_size = Pt(18 if level == 1 else 16 if level == 2 else 14 if level == 3 else 12)
        set_run_font(run, font_size=font_size, bold=True)
    # 设置标题段落格式
    heading.paragraph_format.space_before = Pt(12 if level == 1 else 6)
    heading.paragraph_format.space_after = Pt(6 if level == 1 else 3)
    heading.paragraph_format.line_spacing = 1.5
    return heading

def create_paragraph(doc, style=None):
    """创建段落"""
    p = doc.add_paragraph(style=style)
    # 设置段落格式
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    return p

def add_formatted_text(paragraph, text):
    """添加带格式的文本到段落"""
    if not text:
        return
    
    # 分割行内代码、强调等
    parts = []
    current = ''
    in_code = False
    in_bold = False
    in_italic = False
    
    i = 0
    while i < len(text):
        char = text[i]
        
        # 处理行内代码 `...`
        if text[i:i+1] == '`' and not in_code:
            if current:
                parts.append(('text', current))
            current = ''
            in_code = True
            i += 1
            continue
        elif text[i:i+1] == '`' and in_code:
            parts.append(('code', current))
            current = ''
            in_code = False
            i += 1
            continue
        
        # 处理加粗 **...**
        if i + 1 < len(text) and text[i:i+2] == '**' and not in_bold and not in_code:
            if current:
                parts.append(('text', current))
            current = ''
            in_bold = True
            i += 2
            continue
        elif i + 1 < len(text) and text[i:i+2] == '**' and in_bold:
            parts.append(('bold', current))
            current = ''
            in_bold = False
            i += 2
            continue
        
        # 处理斜体 *...*
        if text[i:i+1] == '*' and not in_italic and not in_code and not in_bold:
            if current:
                parts.append(('text', current))
            current = ''
            in_italic = True
            i += 1
            continue
        elif text[i:i+1] == '*' and in_italic:
            parts.append(('italic', current))
            current = ''
            in_italic = False
            i += 1
            continue
        
        current += char
        i += 1
    
    if current:
        if in_code:
            parts.append(('code', current))
        elif in_bold:
            parts.append(('bold', current))
        elif in_italic:
            parts.append(('italic', current))
        else:
            parts.append(('text', current))
    
    for part_type, content in parts:
        run = paragraph.add_run(content)
        if part_type == 'code':
            set_run_font(run, font_size=Pt(9), color=RGBColor(0, 100, 0))
            run.font.name = 'Consolas'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        elif part_type == 'bold':
            set_run_font(run, bold=True)
        elif part_type == 'italic':
            set_run_font(run, italic=True)
        else:
            set_run_font(run)

def create_table(doc, data, headers=None):
    """创建表格"""
    if not data:
        return None

    rows = len(data)
    cols = len(data[0]) if data else 0

    if rows == 0 or cols == 0:
        return None

    table = doc.add_table(rows=rows + (1 if headers else 0), cols=cols)
    table.style = 'Light Grid Accent 1'

    # 设置表格对齐
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT

    start_idx = 0

    # 设置表头
    if headers:
        header_row = table.rows[0]
        for i, header in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = header
            # 表头背景色
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'D9E2F3')
            cell._element.get_or_add_tcPr().append(shading_elm)

            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    run.font.size = Pt(10.5)
        start_idx = 1

    # 填充数据
    for i, row_data in enumerate(data):
        row = table.rows[i + start_idx]
        for j, cell_data in enumerate(row_data):
            if j < cols:
                cell = row.cells[j]
                cell.text = str(cell_data) if cell_data else ''
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing = 1.5
                    for run in paragraph.runs:
                        run.font.name = '微软雅黑'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                        run.font.size = Pt(10)

    return table

def parse_markdown_to_docx(md_path, output_path):
    """解析 Markdown 文件并生成 DOCX"""

    md_content = Path(md_path).read_text(encoding='utf-8')

    # 清理 LLM 输出的 <think> 标签
    import re
    md_content = re.sub(r'<think>.*?</think>', '', md_content, flags=re.DOTALL)
    md_content = re.sub(r'<thinking>.*?</thinking>', '', md_content, flags=re.DOTALL)

    lines = md_content.split('\n')
    
    doc = Document()
    
    # 设置全局样式
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.font.size = Pt(10.5)
    
    # 提取文档标题（从第一个 # 标题）
    doc_title = '系统功能需求说明书'
    for line in lines[:50]:
        if line.strip().startswith('# '):
            doc_title = line.strip()[2:].strip()
            break
    
    # 解析正文
    in_table = False
    table_data = []
    table_headers = None
    in_code_block = False
    code_block_content = []
    list_indent = []
    
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # 代码块
        if stripped.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_block_content = []
            else:
                # 输出代码块
                p = doc.add_paragraph()
                p.style = 'No Spacing'
                p.paragraph_format.left_indent = Cm(0.5)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                # 添加浅灰色背景
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'F5F5F5')
                p._element.get_or_add_pPr().append(shading_elm)

                for code_line in code_block_content:
                    code_run = p.add_run(code_line + '\n')
                    code_run.font.name = 'Consolas'
                    code_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    code_run.font.size = Pt(9)
                    code_run.font.color.rgb = RGBColor(0, 100, 0)
                in_code_block = False
                code_block_content = []
            i += 1
            continue
        
        if in_code_block:
            code_block_content.append(line)
            i += 1
            continue
        
        # 空行
        if not stripped:
            i += 1
            continue
        
        # 标题处理
        if stripped.startswith('# '):
            create_heading(doc, stripped[2:], level=1)
            i += 1
        elif stripped.startswith('## '):
            create_heading(doc, stripped[3:], level=2)
            i += 1
        elif stripped.startswith('### '):
            create_heading(doc, stripped[4:], level=3)
            i += 1
        elif stripped.startswith('#### '):
            create_heading(doc, stripped[5:], level=4)
            i += 1
        # 表格处理
        elif stripped.startswith('|') and '|' in stripped:
            cells = [cell.strip() for cell in stripped.split('|')[1:-1]]
            
            if all(c.startswith('---') or c == '' for c in cells):
                if len(table_data) > 0:
                    table_headers = table_data[0]
                    table_data = table_data[1:]
                i += 1
                continue
            
            table_data.append(cells)
            
            # 检查是否继续是表格行
            next_i = i + 1
            is_table_end = True
            if next_i < len(lines):
                next_stripped = lines[next_i].strip()
                if next_stripped.startswith('|') and '|' in next_stripped:
                    is_table_end = False
            
            if is_table_end and table_data:
                hdrs = table_headers if table_headers else None
                create_table(doc, table_data, hdrs)
                table_data = []
                table_headers = None
            i += 1
        # 列表处理
        elif stripped.startswith('- ') or stripped.startswith('* '):
            content = stripped[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Cm(0.5)
            add_formatted_text(p, content)
            i += 1
        elif stripped.startswith('1. ') or stripped.startswith('1)'):
            # 有序列表
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Cm(0.5)
            # 移除序号，使用纯文本
            match = re.match(r'^\d+[\.\)]\s*', stripped)
            if match:
                content = stripped[match.end():].strip()
            else:
                content = stripped
            add_formatted_text(p, content)
            i += 1
        # 引用块
        elif stripped.startswith('> '):
            content = stripped[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run('│ ' + content)
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            run.font.color.rgb = RGBColor(100, 100, 100)
            run.font.italic = True
            i += 1
        # 普通段落
        else:
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.first_line_indent = Cm(0.75)  # 首行缩进
            add_formatted_text(p, stripped)
            i += 1
    
    doc.save(output_path)
    print(f"DOCX 文件已生成：{output_path}")

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("用法：python generate_docx.py <input.md> <output.docx>")
        sys.exit(1)
    
    input_path = sys.argv[1]
    output_path = sys.argv[2]
    
    if not Path(input_path).exists():
        print(f"错误：输入文件不存在：{input_path}")
        sys.exit(1)
    
    parse_markdown_to_docx(input_path, output_path)